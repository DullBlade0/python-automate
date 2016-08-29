import docx

def addLines(guestName):
    

doc = docx.Document()
list = open('guests.txt')
letterBody = ['It would be a pleasure to have the company of','' ,'at 11010 Memory Lane on the Evening of', 'April 1st', 'at 7 o\'clock']
styleList = []

for item in list:    
    letterBody[1] = item
    for i, line in enumerate(letterBody):
    doc.add_paragraph(line)
    doc.paragraphs[i].style = 
